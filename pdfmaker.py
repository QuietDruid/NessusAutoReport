import csv
import time

import pandas as pd
import os
import subprocess
import glob
from docx import Document
from pathlib import Path
import time
from datetime import date, timedelta
        
def compileresults(csv_list):
    
    temp = dict()

    for entry in csv_list:
        output = pd.read_csv(entry)
        sorted = output.sort_values(by=["CVSS v2.0 Base Score"], ascending=False)
    
        for x in sorted.values:
            if x[2] >= 5:
                if x[2] in temp:
                    templist = temp.get(x[2])
                    temptuple = (x[4], x[7])
                    templist.add(temptuple)
                    temp.update({x[2]: templist})
                else:
                    templist = set()
                    temptuple = (x[4], x[7])
                    templist.add((temptuple))
                    temp.update({x[2]: templist})
    print(temp)
    
    return temp

def pdfwriter(pathtofile, csv_list):
    docu = Document('rep-template-v0.docx')
    timeofreport = date.today()
    scanstart = (date.today() - timedelta(days=30)).strftime("%b %d")
    timeofreport = timeofreport.strftime("%b %d")
    low = 0
    med = 0
    high = 0
    crit = 0
    maxlist = set()
    averagehost = 0
    maxhost = 0
    for entry in csv_list:
        output = pd.read_csv(entry)
        try:
            risks = output['Risk'].values.tolist()
        except:
            print("CSV doesn't match Nessus formatting. Are you sure that this is the correct file?")
            exit()
        low += risks.count('Low')
        med += risks.count('Medium')
        high += risks.count('High')
        crit += risks.count('Critical')
        for x in output.values:
            maxlist.add(x[4])
        if len(maxlist) > maxhost:
            maxhost = len(maxlist)
        averagehost += len(maxlist)
    low = low  // len(csv_list)
    med = med  // len(csv_list)
    high = high  // len(csv_list)
    crit = crit  // len(csv_list)

    pathtofile = pathtofile + "/latestreport.docx"
    print("\nPath To Latest Report in Docx: " + pathtofile)
    docu.save(pathtofile)
    docu.tables[0].cell(0, 1).text = str(maxhost)
    docu.tables[0].cell(0, 3).text = str(averagehost // len(csv_list))
    docu.tables[0].cell(1, 1).text = scanstart
    docu.tables[0].cell(1, 3).text = timeofreport

    docu.tables[1].cell(0, 1).text = str(crit)
    docu.tables[1].cell(0, 1).paragraphs[0].alignment = 1

    docu.tables[1].cell(1, 1).text = str(high)
    docu.tables[1].cell(1, 1).paragraphs[0].alignment = 1

    docu.tables[1].cell(2, 1).text = str(med)
    docu.tables[1].cell(2, 1).paragraphs[0].alignment = 1

    docu.tables[1].cell(3, 1).text = str(low)
    docu.tables[1].cell(3, 1).paragraphs[0].alignment = 1

    docu.tables[1].cell(4, 1).text = str(crit + high + med + low)
    docu.tables[1].cell(4, 1).paragraphs[0].alignment = 1
    docu.save(pathtofile)


def vulnwriter(pathtofile, csv_list):
    vulnlist = compileresults(csv_list)
    docu = Document(pathtofile + "/latestreport.docx")
    count = 1
    print("\nVulnerabilities Being Written to Report:")
    
    for vulnscore in vulnlist.keys():
        for entry in vulnlist.get(vulnscore):
            if len(docu.tables[2].rows) == count:
                docu.tables[2].add_row()
                docu.tables[2].style = 'Table Grid'
            
            docu.tables[2].cell(count, 0).text = entry[0]
            docu.tables[2].cell(count, 0).paragraphs[0].alignment = 1

            docu.tables[2].cell(count, 1).text = entry[1]
            docu.tables[2].cell(count, 1).paragraphs[0].alignment = 1

            docu.tables[2].cell(count, 2).text = str(vulnscore)
            docu.tables[2].cell(count, 2).paragraphs[0].alignment = 1

            count += 1
            print(str(vulnscore) + " " + entry[0])
    docu.save(pathtofile + "/latestreport.docx")
    print()


def callpdfmaker():
    #Gets path to downloads folder to get the csv files
    path = str(Path.home() / "Downloads" / "csvs")
    #Following gets date 30 days from now and creates a list of csv files names from within 30 days
    scanstart = (date.today() - timedelta(days=30))
    recent_csv = []
    csv_files = glob.glob(os.path.join(path, "*.csv"))
    

    for file_path in csv_files:
        mtime = date.fromtimestamp(os.path.getmtime(file_path))
        if mtime >= scanstart:
            recent_csv.append(file_path)
    print(recent_csv)
    #pdf writer initializes the working doc files and adds some preset data at the start
    pdfwriter(path, recent_csv) 
    #vulnwriter scours csvs and compiles vulnerabilities
    vulnwriter(path, recent_csv)

    docxtemp = path + "/latestreport.docx"
    outpdf = path + "/latestreport.pdf"
    #time.sleep(10)
    print("Temp Files For Editing:")
    print(docxtemp)
    print(outpdf + "\nConverting Docx to PDF")
    os.chdir("/home/starphyre/Downloads/csvs")
    os.system("libreoffice --headless --convert-to pdf latestreport.docx")



if __name__ == '__main__':
    #This is redundant but I wanted to allow it to run by itself
    callpdfmaker()
